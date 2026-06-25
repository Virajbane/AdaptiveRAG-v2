"""
Unit tests for document parsers (TXT, CSV, PDF, DOCX).

TXT and CSV tests create temp files inline - no extra dependencies.
PDF and DOCX tests build minimal real files using the same libraries
the parsers themselves use (PyMuPDF / python-docx), so we're testing
against real file structures rather than guessing at byte layouts.
"""

import os
import tempfile
import pytest
import fitz
from docx import Document as DocxDocument

from app.services.document.parser import (
    TXTParser,
    CSVParser,
    PDFParser,
    DOCXParser,
    DocumentParser,
)


# ── TXT Parser ──────────────────────────────────────────────────

def test_txt_parser_reads_content():
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".txt", delete=False, encoding="utf-8"
    ) as f:
        f.write("Hello, this is a test document.")
        path = f.name

    try:
        result = TXTParser.parse(path)
        assert "Hello, this is a test document." in result
    finally:
        os.unlink(path)


def test_txt_parser_missing_file_raises():
    with pytest.raises(ValueError):
        TXTParser.parse("/nonexistent/path/file.txt")


# ── CSV Parser ──────────────────────────────────────────────────

def test_csv_parser_includes_headers():
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".csv", delete=False, encoding="utf-8", newline=""
    ) as f:
        f.write("name,age\nAlice,30\nBob,25\n")
        path = f.name

    try:
        result = CSVParser.parse(path)
        assert "HEADERS: name | age" in result
        assert "Alice | 30" in result
        assert "Bob | 25" in result
    finally:
        os.unlink(path)


def test_csv_parser_missing_file_raises():
    with pytest.raises(ValueError):
        CSVParser.parse("/nonexistent/path/file.csv")


# ── PDF Parser ──────────────────────────────────────────────────

def test_pdf_parser_extracts_text():
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        path = f.name

    try:
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "This is page one content.")
        doc.save(path)
        doc.close()

        result = PDFParser.parse(path)
        assert "This is page one content." in result
        assert "Page 1" in result
    finally:
        os.unlink(path)


def test_pdf_parser_missing_file_raises():
    with pytest.raises(ValueError):
        PDFParser.parse("/nonexistent/path/file.pdf")


# ── DOCX Parser ─────────────────────────────────────────────────

def test_docx_parser_extracts_paragraphs():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = DocxDocument()
        doc.add_paragraph("This is a test paragraph.")
        doc.save(path)

        result = DOCXParser.parse(path)
        assert "This is a test paragraph." in result
    finally:
        os.unlink(path)


def test_docx_parser_extracts_tables():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        doc = DocxDocument()
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Score"
        doc.save(path)

        result = DOCXParser.parse(path)
        assert "[TABLE]" in result
        assert "Name | Score" in result
    finally:
        os.unlink(path)


def test_docx_parser_missing_file_raises():
    with pytest.raises(ValueError):
        DOCXParser.parse("/nonexistent/path/file.docx")


# ── DocumentParser dispatcher ───────────────────────────────────

def test_document_parser_dispatches_to_txt():
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".txt", delete=False, encoding="utf-8"
    ) as f:
        f.write("dispatcher test content")
        path = f.name

    try:
        result = DocumentParser.parse(path, "txt")
        assert "dispatcher test content" in result
    finally:
        os.unlink(path)


def test_document_parser_unsupported_type_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        DocumentParser.parse("/some/path.xyz", "xyz")


def test_document_parser_does_not_support_markdown_yet():
    """
    KNOWN GAP: the project's original requirements list Markdown
    (.md) as a supported upload format, but DocumentParser.PARSERS
    has no 'md' entry. This test documents that gap explicitly so
    it's a visible, tracked omission rather than a silent one.

    If markdown support is added later, this test should be removed
    or flipped to assert successful parsing instead.
    """
    with pytest.raises(ValueError, match="Unsupported file type"):
        DocumentParser.parse("/some/path.md", "md")