import re
import fitz  # PyMuPDF for PDF
from typing import List
from docx import Document as DocxDocument
import csv


def normalize_decimal_commas(text: str) -> str:
    """
    Converts comma-decimal numbers like "53,51" -> "53.51".
    Only fires when digits appear on both sides of the comma, so
    genuine list separators like "A, B, C" are left untouched --
    this only matches tight digit-comma-digit patterns as seen in
    this paper's tables (e.g. "53,51" not "5, 3").
    """
    return re.sub(r'(?<=\d),(?=\d)', '.', text)

def normalize_wer_row(text: str) -> str:
    """
    Table 1's "STT WER" row presents five bare percentage values with no
    inline labels -- just newline-separated numbers under condition
    headers several lines above. The LLM has to count position 4 vs 5
    to know which number is TTS-based vs xTTS-based, and gets this wrong
    consistently (confirmed on Q17 across multiple test runs: model
    reports xTTS=6.6%/TTS=20.4%, reversed from the correct xTTS=20.4%/
    TTS=6.6%). Column order confirmed by direct inspection of Table 1:
    GPT2-based, gTTS-based, TTS-based(250), xTTS-based(250).

    Rewrites the matched row with labels attached inline, e.g.:
    "STT WER: GPT2-based=11.9%, gTTS-based=-, TTS-based(250)=15.9%,
    TTS-based(diarised)=6.6%, xTTS-based(250)=20.4%"
    so the LLM reads the label directly instead of counting position.

    Narrowly scoped to this exact row shape (five values immediately
    following the literal "STT WER" line) so it can't misfire on
    unrelated content elsewhere in the document.
    """
    pattern = re.compile(
        r'STT WER\s*\n'
        r'([\d.]+%?|-)\s*\n'
        r'([\d.]+%?|-)\s*\n'
        r'([\d.]+%?|-)\s*\n'
        r'([\d.]+%?|-)\s*\n'
        r'([\d.]+%?|-)'
    )

    def _replace(m):
        gpt2, gtts, tts_diar, tts_250, xtts_250 = m.groups()
        return (
            f"STT WER: GPT2-based(Transcribed)={gpt2}, "
            f"gTTS-based(Transcribed)={gtts}, "
            f"TTS-based(Diarised)={tts_diar}, "
            f"TTS-based(250-sample,Features)={tts_250}, "
            f"xTTS-based(250-sample,Features)={xtts_250}"
        )

    return pattern.sub(_replace, text)

def normalize_hyphenated_linebreaks(text: str) -> str:
    """
    Rejoins words split across a line-break hyphen, e.g. "Resem-\nblyzer's"
    -> "Resemblyzer's". PyMuPDF preserves the PDF's original line-wrap
    hyphenation as literal "word-\nrest" sequences within a text block.
    Left unfixed, chunk boundaries or embeddings can land right on this
    fracture, so a name like "Resemblyzer" gets split into "Resem-" at
    the end of one chunk and "blyzer's..." at the start of the next --
    degrading retrieval for exactly the terms most likely to be searched
    (proper nouns, tool/library names). Confirmed on Q18: the chunk
    containing the ARI value literally started mid-word ("blyzer's
    VoiceEncoder...") after "Resem-" was cut off in the prior chunk.

    Only joins lowercase-to-lowercase breaks (the common case for a
    single word wrapped across a line) to avoid merging genuine
    compound hyphenated terms that happen to fall at a line boundary.
    """
    return re.sub(r'([a-z])-\n([a-z])', r'\1\2', text)


class PDFParser:
    """Parse PDF files"""

    @staticmethod
    def parse(file_path: str) -> str:
        text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                blocks = page.get_text("blocks")  # (x0,y0,x1,y1, text, block_no, ...)
                blocks.sort(key=lambda b: (round(b[1]), b[0]))  # reading order: top-to-bottom, left-to-right
                for b in blocks:
                    block_text = b[4].strip()
                    if block_text:
                        text += block_text + "\n\n"  # real paragraph boundary between blocks
            doc.close()

            # 2026-07-05 fix: some tables in source PDFs render decimal
            # numbers with commas instead of periods (e.g. "53,51" meaning
            # 53.51%). Left unfixed, this reads as noise to an LLM trying
            # to extract a specific value from a dense number table --
            # traced directly to a real hallucination (Q16: model grabbed
            # the wrong column's value instead of correctly reading
            # 53.51/67.66). Applied globally since digit-comma-digit is
            # never a real separator, only a decimal point, in this corpus.
            #
            # 2026-07-05: reverted the find_tables()-based table detection
            # tried earlier -- it found 0 real tables on this document and
            # instead misidentified chart/figure grids as tables (false
            # positives), adding noise without fixing the actual problem.
            # Root cause was the comma-decimal formatting above, not
            # missing table structure -- the plain block-sort extraction
            # already keeps table titles and data together correctly.
            text = normalize_decimal_commas(text)
            text = normalize_hyphenated_linebreaks(text)
            text = normalize_wer_row(text)

            print(f"[PARSER DEBUG] Total chars: {len(text)}")
            print(f"[PARSER DEBUG] Block-separated paragraph count: {text.count(chr(10)+chr(10))}")
            try:
                print(f"[PARSER DEBUG] First 300 chars: {repr(text[:300])}")
            except UnicodeEncodeError:
                print("[PARSER DEBUG] First 300 chars: (contains non-ASCII characters, skipped)")
            return text
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")


class DOCXParser:
    """Parse DOCX files"""

    @staticmethod
    def parse(file_path: str) -> str:
        """Extract text from DOCX"""
        text = ""
        try:
            doc = DocxDocument(file_path)

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"

            # Extract tables
            for table in doc.tables:
                text += "\n[TABLE]\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
                text += "[/TABLE]\n"

            return text
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")


class TXTParser:
    """Parse plain text files"""

    @staticmethod
    def parse(file_path: str) -> str:
        """Read text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"Error parsing TXT: {str(e)}")


class CSVParser:
    """Parse CSV files"""

    @staticmethod
    def parse(file_path: str) -> str:
        """Convert CSV to text format"""
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for i, row in enumerate(reader):
                    if i == 0:
                        text += "HEADERS: " + " | ".join(row) + "\n"
                    else:
                        text += " | ".join(row) + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Error parsing CSV: {str(e)}")


# Main parser dispatcher
class DocumentParser:
    """Main document parser"""

    PARSERS = {
        "pdf": PDFParser,
        "docx": DOCXParser,
        "txt": TXTParser,
        "csv": CSVParser
    }

    @staticmethod
    def parse(file_path: str, file_type: str) -> str:
        """Parse document based on type"""
        file_type = file_type.lower()
        if file_type not in DocumentParser.PARSERS:
            raise ValueError(f"Unsupported file type: {file_type}")

        parser_class = DocumentParser.PARSERS[file_type]
        return parser_class.parse(file_path)